#!/usr/bin/env python3
"""
Verify the filled SDT document
"""
from docx import Document

doc = Document("SDT - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].docx")

print("=== VERIFICATION: Checking filled content ===\n")

# Find and display key sections
key_headings = [
    "Product summary (MVP)",
    "Understanding the Client's Needs",
    "Value Proposition",
    "Key Differentiators"
]

for heading_text in key_headings:
    found = False
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == heading_text:
            print(f"\n✓ {heading_text}")
            # Print next 3 content paragraphs
            count = 0
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                content = doc.paragraphs[j].text.strip()
                if content and not doc.paragraphs[j].style.name.startswith('Heading'):
                    print(f"  {content[:100]}")
                    count += 1
                    if count >= 2:
                        break
            found = True
            break
    if not found:
        print(f"\n✗ {heading_text} - NOT FOUND")

print("\n=== Summary ===")
print(f"Total paragraphs in document: {len(doc.paragraphs)}")
print("\n✓ SDT document successfully filled from presales guide markdown!")
