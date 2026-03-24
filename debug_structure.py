#!/usr/bin/env python3
"""
Debug script to see exact heading structure and content
"""
from docx import Document

doc = Document("templates/NEW presales_template_sdt.docx")

print("=== TEMPLATE STRUCTURE (Headings and immediate content) ===\n")

for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading'):
        print(f"\nPara {i}: [{para.style.name}] {para.text}")
        # Show next 2 content lines
        for j in range(i+1, min(i+3, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            if next_para.style.name.startswith('Heading'):
                break
            text = next_para.text.strip()
            if text and not text.startswith('['):
                print(f"    Content: {text[:80]}")
