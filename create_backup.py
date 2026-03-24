#!/usr/bin/env python3
"""
Create a clean copy of the SDT document with simpler name
"""
from docx import Document
import shutil
import os

source = 'SDT - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].docx'
backup = 'SDT_DBMS_Oracle_Azure_FILLED.docx'

print(f"Creating backup copy: {backup}")

try:
    # Load the document
    doc = Document(source)
    print(f"✓ Loaded source document")
    
    # Save with simpler name
    doc.save(backup)
    print(f"✓ Saved backup: {backup}")
    
    # Verify the backup
    doc_check = Document(backup)
    sdts = doc_check.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt')
    print(f"✓ Backup verified ({len(sdts)} SDT fields, {len(doc_check.paragraphs)} paragraphs)")
    
    # Also try copying to root
    dest_root = f"C:\\Users\\koengo\\Desktop\\{backup}"
    shutil.copy(backup, dest_root)
    print(f"✓ Also copied to Desktop: {dest_root}")
    
except Exception as e:
    print(f"✗ Error: {e}")
    import traceback
    traceback.print_exc()
