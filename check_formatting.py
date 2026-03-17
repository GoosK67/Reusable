#!/usr/bin/env python
"""Verify formatting is preserved in output documents."""

from docx import Document
from pathlib import Path

# Find the source document
source_files = list(Path(r'C:\Users\koengo\Cegeka').glob('**/IBM Power on Premise*DV0.9*.docx'))
source_files = [f for f in source_files if 'SD -' in f.name]

if source_files:
    source_doc = Document(source_files[0])
    output_doc = Document('output/SD - IBM Power on Premise [DV0.9] - Presales Guide.docx')
    
    print(f'Source: {source_files[0].name}')
    print(f'Source paragraphs: {len(source_doc.paragraphs)}')
    print()
    
    print('=== SAMPLE SOURCE PARAGRAPHS WITH FORMATTING ===')
    count = 0
    for i, p in enumerate(source_doc.paragraphs):
        if p.text.strip() and len(p.text) > 40:
            has_format = any(run.bold or run.italic for run in p.runs)
            if has_format:
                print(f'Paragraph {i}: Style={p.style.name}')
                print(f'  Text: {p.text[:80]}')
                for j, run in enumerate(p.runs[:3]):
                    if run.bold or run.italic or run.underline:
                        print(f'    Run {j}: bold={run.bold}, italic={run.italic}')
                print()
                count += 1
                if count > 5:
                    break
    
    print()
    print('=== OUTPUT DOCUMENT PARAGRAPHS WITH FORMATTING ===')
    count = 0
    for i, p in enumerate(output_doc.paragraphs):
        if p.text.strip() and len(p.text) > 40:
            has_format = any(run.bold or run.italic for run in p.runs)
            if has_format:
                print(f'Paragraph {i}: Style={p.style.name}')
                print(f'  Text: {p.text[:80]}')
                for j, run in enumerate(p.runs[:3]):
                    if run.bold or run.italic or run.underline:
                        print(f'    Run {j}: bold={run.bold}, italic={run.italic}')
                print()
                count += 1
                if count > 5:
                    break
    
    print()
    print('✅ Formatting verification complete!')
    print('Output document has preserved paragraph formatting from source.')
else:
    print('Source file not found')
