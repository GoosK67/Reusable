#!/usr/bin/env python3
"""
Inspect the SDT template structure for form fields and content controls
"""
from docx import Document
from docx.oxml import parse_xml
import xml.etree.ElementTree as ET

doc = Document('templates/presales_template_sdt_v2.docx')

print("=== TEMPLATE INSPECTION ===\n")

# Check for form fields (legacy)
print("1. Checking for form fields...")
form_fields = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldSimple')
print(f"   Found {len(form_fields)} legacy form fields")

# Check for content controls (structured)
print("\n2. Checking for content controls (SDT)...")
sdts = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt')
print(f"   Found {len(sdts)} content controls (SDT elements)")

if sdts:
    print("\n   Content Control Tags:")
    for i, sdt in enumerate(sdts[:20]):  # Show first 20
        # Try to get the tag
        sdtPr = sdt.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr')
        if sdtPr is not None:
            tag_elem = sdtPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag')
            if tag_elem is not None:
                tag_val = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                print(f"   [{i}] Tag: {tag_val}")
            else:
                print(f"   [{i}] No tag attribute found")

# Check for bookmarks
print("\n3. Checking for bookmarks...")
bookmarks = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart')
print(f"   Found {len(bookmarks)} bookmarks")
if bookmarks:
    for i, bm in enumerate(bookmarks[:10]):
        name = bm.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
        print(f"   [{i}] {name}")

# Show structure
print("\n4. Document paragraphs and structure:")
for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()[:60]
    if text:
        print(f"   Para {i}: {text}")

print("\n5. Checking tables for content controls...")
for table_idx, table in enumerate(doc.tables[:3]):
    print(f"\n   Table {table_idx}:")
    for row_idx, row in enumerate(table.rows[:3]):
        for cell_idx, cell in enumerate(row.cells):
            # Check for SDTs in cells
            cell_sdts = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt')
            if cell_sdts:
                print(f"     Row {row_idx}, Cell {cell_idx}: {len(cell_sdts)} SDT(s)")
                for sdt in cell_sdts:
                    sdtPr = sdt.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr')
                    if sdtPr is not None:
                        tag_elem = sdtPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag')
                        if tag_elem is not None:
                            tag_val = tag_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            print(f"       Tag: {tag_val}")
