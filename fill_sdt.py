#!/usr/bin/env python3
"""
Fill the SDT template with presales guide information
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Read the presales markdown
with open('presales/Presales Guide - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Read the template
doc = Document('templates/NEW presales_template_sdt.docx')

# Extract sections from markdown
sections = {}
current_section = None
current_subsection = None
current_content = []

lines = md_content.split('\n')
for line in lines:
    # Check for main section (## format)
    if line.startswith('## '):
        # Save previous content
        if current_section and current_content:
            if current_subsection:
                if not isinstance(sections[current_section], dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            else:
                if current_section not in sections:
                    sections[current_section] = '\n'.join(current_content).strip()
        
        current_section = line[3:].strip()
        current_subsection = None
        current_content = []
    # Check for subsection (### format)
    elif line.startswith('### '):
        # Save previous content
        if current_content:
            if current_subsection:
                if not isinstance(sections[current_section], dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            elif current_section:
                if current_section not in sections:
                    sections[current_section] = '\n'.join(current_content).strip()
        
        # Initialize subsection dict if needed
        if current_section and not isinstance(sections.get(current_section, {}), dict):
            old_content = sections.get(current_section, '')
            sections[current_section] = {'_main': old_content} if old_content else {}
        
        current_subsection = line[4:].strip()
        current_content = []
    elif current_section or current_subsection:
        if line.strip():
            current_content.append(line)

# Capture last section
if current_section and current_content:
    if current_subsection:
        if not isinstance(sections.get(current_section, {}), dict):
            sections[current_section] = {}
        sections[current_section][current_subsection] = '\n'.join(current_content).strip()
    else:
        sections[current_section] = '\n'.join(current_content).strip()

print("Sections extracted from markdown:")
for section in sections.keys():
    print(f"  - {section}")

# Now fill the template - let's first output the template structure
print("\n\nTemplate structure (all headings):")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and para.style.name.startswith('Heading'):
        print(f"Para {i}: [{para.style.name}] {text[:100]}")

# Save template info to file for inspection
with open('template_structure.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"{i}: [{para.style.name}] {para.text}\n")
