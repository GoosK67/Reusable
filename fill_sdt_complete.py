#!/usr/bin/env python3
"""
Fill the SDT template with presales guide information
Maps markdown sections to template sections
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Read the presales markdown
with open('presales/Presales Guide - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Extract sections from markdown
sections = {}
current_section = None
current_subsection = None
current_content = []

lines = md_content.split('\n')
for line in lines:
    if line.startswith('## '):
        if current_section and current_content:
            if current_subsection:
                if not isinstance(sections.get(current_section), dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            else:
                sections[current_section] = '\n'.join(current_content).strip()
        
        current_section = line[3:].strip()
        current_subsection = None
        current_content = []
    elif line.startswith('### '):
        if current_content:
            if current_subsection:
                if not isinstance(sections.get(current_section), dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            elif current_section:
                if current_section not in sections:
                    sections[current_section] = '\n'.join(current_content).strip()
        
        if current_section and not isinstance(sections.get(current_section), dict):
            old_content = sections.get(current_section, '')
            sections[current_section] = {'_main': old_content} if old_content else {}
        
        current_subsection = line[4:].strip()
        current_content = []
    elif current_section or current_subsection:
        if line.strip():
            current_content.append(line)

if current_section and current_content:
    if current_subsection:
        if not isinstance(sections.get(current_section), dict):
            sections[current_section] = {}
        sections[current_section][current_subsection] = '\n'.join(current_content).strip()
    else:
        sections[current_section] = '\n'.join(current_content).strip()

print("✓ Extracted sections from markdown:")
for section in sections.keys():
    print(f"  - {section}")

# Load template
doc = Document('templates/NEW presales_template_sdt.docx')
print("\n✓ Loaded template")

# Helper function to clear a section and add new content after a heading
def fill_section_after_heading(doc, heading_para_index, content_text):
    """Replace content after a heading with new text, preserving formatting"""
    if not content_text or not content_text.strip():
        return
    
    # Find where this section ends (next heading or end of document)
    end_index = len(doc.paragraphs)
    for i in range(heading_para_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            end_index = i
            break
    
    # Remove old paragraphs between heading and next heading
    for i in range(end_index - 1, heading_para_index, -1):
        p = doc.paragraphs[i]
        # Only remove if it's not a table or the heading itself
        if p != doc.paragraphs[heading_para_index]:
            pPr = p._element.getparent()
            pPr.remove(p._element)
    
    # Add new content after the heading
    heading_para = doc.paragraphs[heading_para_index]
    
    # Parse the content and add as bullets or paragraphs
    content_lines = content_text.strip().split('\n')
    
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if it's a bullet point
        if line.startswith('- '):
            p = heading_para.insert_paragraph_before(line[2:])
            p.style = 'List Bullet'
        else:
            p = heading_para.insert_paragraph_before(line)
            p.style = 'Normal'

# Mapping of markdown sections to template heading indices
# Based on the structure we found earlier
mapping = {
    '1. Product Summary': 44,
    '2. Understanding the Client\'s Needs': 48,
    '3. Product Description': (51, {
        '3.1 Architectural Description': 54,
        '3.2 Key Features & Functionalities': 59,
        '3.3 Scope / Out-of-Scope': 61,
        '3.4 Requirements & Prerequisites': 63,
    }),
    '4. Value Proposition': 66,
    '5. Key Differentiators': 75,
    '6. Transition & Transformation': 83,
    '7. Client Responsibilities': 103,
    '8. Operational Support': 106,
    '9. Terms & Conditions': 111,
    '10. SLA & KPI Management': 119,
    '11. Pricing Elements': 130,
}

print("\nFilling template sections:")

# Get actual paragraph indices by looking for headings in the document
para_by_heading = {}
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        para_by_heading[text] = i

# Fill each section
for md_section, template_info in mapping.items():
    if md_section not in sections:
        print(f"  ⚠ {md_section} not found in markdown")
        continue
    
    # Handle nested subsections
    if isinstance(template_info, tuple):
        main_heading_idx, subsections = template_info
        main_content = sections[md_section]
        
        # If main section is a dict (has subsections), get the _main content
        if isinstance(main_content, dict):
            main_text = main_content.get('_main', '')
        else:
            main_text = main_content
        
        if main_text:
            print(f"  ✓ Filling {md_section}")
            fill_section_after_heading(doc, main_heading_idx, main_text)
        
        # Fill subsections
        if isinstance(main_content, dict):
            for sub_key, template_sub_idx in subsections.items():
                if sub_key in main_content:
                    print(f"  ✓ Filling {sub_key}")
                    fill_section_after_heading(doc, template_sub_idx, main_content[sub_key])
    else:
        template_idx = template_info
        content = sections[md_section]
        
        if isinstance(content, dict):
            content_text = content.get('_main', '')
        else:
            content_text = content
        
        if content_text:
            print(f"  ✓ Filling {md_section}")
            fill_section_after_heading(doc, template_idx, content_text)

# Save the filled template
output_filename = 'SDT - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].docx'
doc.save(output_filename)
print(f"\n✓ Saved filled SDT to: {output_filename}")
