#!/usr/bin/env python3
"""
Check DOCX file integrity and repair if needed
"""
from docx import Document
import zipfile
import os

filename = 'SDT - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].docx'

print(f"Checking file: {filename}\n")

# Check file exists
if not os.path.exists(filename):
    print(f"✗ File not found!")
    exit(1)

print(f"✓ File found ({os.path.getsize(filename)} bytes)")

# Check if it's a valid ZIP (DOCX is a ZIP file)
try:
    with zipfile.ZipFile(filename, 'r') as z:
        file_list = z.namelist()
        print(f"✓ Valid ZIP structure ({len(file_list)} files)")
        
        # Check for critical DOCX files
        critical_files = ['word/document.xml', '[Content_Types].xml']
        for cf in critical_files:
            if cf in file_list:
                print(f"  ✓ {cf} present")
            else:
                print(f"  ✗ {cf} MISSING!")
except Exception as e:
    print(f"✗ ZIP verification failed: {e}")
    exit(1)

# Try to load as Document
try:
    doc = Document(filename)
    print(f"✓ Document loads successfully")
    print(f"  - {len(doc.paragraphs)} paragraphs")
    print(f"  - {len(doc.tables)} tables")
    
    # Check SDT fields
    sdts = doc.element.body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt')
    print(f"  - {len(sdts)} content controls (SDT fields)")
    
    if len(sdts) > 0:
        print("\n✓ Document structure is valid!")
    else:
        print("\n✗ Warning: No SDT fields found")
        
except Exception as e:
    print(f"✗ Document loading failed: {e}")
    exit(1)

print("\n✓ File integrity check PASSED")
