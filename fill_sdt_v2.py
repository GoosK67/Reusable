#!/usr/bin/env python3
"""
Fill the SDT template with presales guide - Version 2
More careful approach using direct text replacement
"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re

# Read the presales markdown
with open('presales/Presales Guide - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Extract sections from markdown
sections = {}
current_section = None
current_subsection = None
current_content = []

lines = md_content.split('\n')
for line in lines:
    if line.startswith('## '):
        if current_section and current_content:
            if current_subsection:
                if not isinstance(sections.get(current_section), dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            else:
                sections[current_section] = '\n'.join(current_content).strip()
        
        current_section = line[3:].strip()
        current_subsection = None
        current_content = []
    elif line.startswith('### '):
        if current_content:
            if current_subsection:
                if not isinstance(sections.get(current_section), dict):
                    sections[current_section] = {}
                sections[current_section][current_subsection] = '\n'.join(current_content).strip()
            elif current_section:
                if current_section not in sections:
                    sections[current_section] = '\n'.join(current_content).strip()
        
        if current_section and not isinstance(sections.get(current_section), dict):
            old_content = sections.get(current_section, '')
            sections[current_section] = {'_main': old_content} if old_content else {}
        
        current_subsection = line[4:].strip()
        current_content = []
    elif current_section or current_subsection:
        if line.strip():
            current_content.append(line)

if current_section and current_content:
    if current_subsection:
        if not isinstance(sections.get(current_section), dict):
            sections[current_section] = {}
        sections[current_section][current_subsection] = '\n'.join(current_content).strip()
    else:
        sections[current_section] = '\n'.join(current_content).strip()

print("✓ Extracted sections from markdown")

# Load template
doc = Document('templates/presales_template_sdt_v2.docx')
print("✓ Loaded template (presales_template_sdt_v2.docx)")

# Helper function to find heading by text (fuzzy match)
def find_heading_index(doc, heading_text):
    """Find a heading by text, with fuzzy matching for special characters"""
    for i, para in enumerate(doc.paragraphs):
        # Clean text for comparison
        para_text_clean = para.text.strip().lower()
        heading_clean = heading_text.lower()
        
        # Try exact match first
        if para_text_clean == heading_clean:
            return i
        
        # Try substring match
        if heading_clean in para_text_clean or para_text_clean in heading_clean:
            if para.style.name.startswith('Heading'):
                return i
    
    return None

# Helper to clear content after heading and add new content
def fill_after_heading(doc, heading_index, new_content_text):
    """Fill the section after a heading with new content"""
    if heading_index is None or not new_content_text.strip():
        return
    
    heading_para = doc.paragraphs[heading_index]
    
    # Find the next heading
    next_heading_idx = None
    for i in range(heading_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            next_heading_idx = i
            break
    
    # Determine range to clear (keep description paragraphs, clear old content)
    clear_from = heading_index + 1
    
    # Skip short descriptive paragraphs (usually less than 15 words)
    while clear_from < (next_heading_idx or len(doc.paragraphs)):
        para = doc.paragraphs[clear_from]
        word_count = len(para.text.strip().split())
        # If it's a short description (< 20 words) and not formatted content, skip it
        if word_count < 20 and para.style.name == 'Normal':
            clear_from += 1
        else:
            break
    
    # Remove paragraphs between clear_from and next heading
    if next_heading_idx:
        # Remove in reverse order to maintain indices
        for i in range(next_heading_idx - 1, clear_from - 1, -1):
            p = doc.paragraphs[i]
            pPr = p._element.getparent()
            pPr.remove(p._element)
    
    # Add new content as paragraphs
    content_lines = new_content_text.strip().split('\n')
    for line_idx, line in enumerate(content_lines):
        line = line.rstrip()
        if not line:
            continue
        
        # Determine style based on line format
        if line.startswith('- '):
            # Bullet point
            text = line[2:].strip()
            p = heading_para.insert_paragraph_before(text)
            p.style = 'List Bullet'
        elif line.startswith('**') and '**' in line[2:]:
            # Bold text - extract and add as normal paragraph
            text = line.replace('**', '')
            p = heading_para.insert_paragraph_before(text)
            p.style = 'Normal'
            p.runs[0].bold = True if p.runs else False
        else:
            # Normal paragraph
            p = heading_para.insert_paragraph_before(line)
            p.style = 'Normal'

# Mapping with exact heading names from the template
print("\nFilling template sections:")

mapping = {
    'Product summary (MVP)': sections.get('1. Product Summary', ''),
    "Understanding the Client": sections.get('2. Understanding the Client\'s Needs', ''),  # Fuzzy match for special char
    'Product Description (MVP)': sections.get('3. Product Description', {}),
    'Architectural description': (sections.get('3. Product Description', {}), '3.1 Architectural Description'),
    'Key features and functionalities': (sections.get('3. Product Description', {}), '3.2 Key Features & Functionalities'),
    'Scope / out-of-scope': (sections.get('3. Product Description', {}), '3.3 Scope / Out-of-Scope'),
    'Requirements and Prerequisites': (sections.get('3. Product Description', {}), '3.4 Requirements & Prerequisites'),
    'Value Proposition': sections.get('4. Value Proposition', ''),
    'Key Differentiators': sections.get('5. Key Differentiators', ''),
    'Transition & Transformation (MVP)': sections.get('6. Transition & Transformation', {}),
    'Client responsibilities': sections.get('7. Client Responsibilities', ''),
    'Operational Support': sections.get('8. Operational Support', ''),
    'Terms and Conditions': sections.get('9. Terms & Conditions', ''),
    'SLA & KPI Management': sections.get('10. SLA & KPI Management', ''),
    'Cost/Pricing elements': sections.get('11. Pricing Elements', ''),
}

for template_heading, md_source in mapping.items():
    idx = find_heading_index(doc, template_heading)
    
    if idx is None:
        print(f"  ⚠ {template_heading} - NOT FOUND")
        continue
    
    # Handle dynamic subsection extraction
    content = ''
    if isinstance(md_source, tuple):
        parent_dict, sub_key = md_source
        if isinstance(parent_dict, dict) and sub_key in parent_dict:
            content = parent_dict[sub_key]
    elif isinstance(md_source, dict):
        content = md_source.get('_main', '')
    else:
        content = md_source
    
    if content:
        print(f"  ✓ {template_heading}")
        fill_after_heading(doc, idx, content)
    else:
        print(f"  ○ {template_heading} (no content to fill)")

# Save the filled template
output_filename = 'SDT - Cegeka DBMS for Oracle on Azure [PRD.0.8.001][PV1.0][DV1.0].docx'
doc.save(output_filename)
print(f"\n✓ Saved filled SDT to: {output_filename}")
