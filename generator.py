from docx import Document
import re


SECTION_TO_FIELD = {
    "product summary": "ProductSummary",
    "product summary (mvp)": "ProductSummary",
    "value proposition": "ValueProposition",
    "product description": "ProductDescription",
    "product description (mvp)": "ProductDescription",
    "architectural description": "ProductDescription",
    "key features & functionalities": "ProductDescription",
    "key features and functionalities": "ProductDescription",
    "requirements & prerequisites": "Requirements",
    "requirements and prerequisites": "Requirements",
    "scope / out of scope": "Scope",
    "scope out-of-scope": "Scope",
    "scope / out-of-scope": "Scope",
    "sla": "SLA",
    "sla & kpi management": "SLA",
    "operational support": "OperationalSupport",
}


def _normalize_heading(text):
    normalized = " ".join(str(text or "").strip().lower().split())
    # Support numbered template headings like "2. Product Summary" or "4.3 Scope / Out-of-Scope".
    normalized = re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", normalized)
    return normalized


def _is_heading_paragraph(paragraph):
    return paragraph.style and paragraph.style.name.startswith("Heading")


def _extract_text(field_value):
    """Extract plain text from either a string or {text, tables} dict."""
    if isinstance(field_value, dict):
        return str(field_value.get("text", "") or "")
    return str(field_value or "")


def copy_paragraph_with_formatting(source_paragraph, target_doc):
    """Copy a paragraph with all formatting from source to target."""
    import copy
    new_p_element = copy.deepcopy(source_paragraph._element)
    target_doc.element.body.append(new_p_element)


def copy_table_with_formatting(source_table, target_doc):
    """Copy a table with all formatting from source to target."""
    import copy
    new_tbl_element = copy.deepcopy(source_table._element)
    target_doc.element.body.append(new_tbl_element)


def insert_doc_tables(doc, table_payloads):
    """Insert actual DOCX table objects into the document.
    
    Args:
        doc: Document object
        table_payloads: List of table row lists, e.g. [[['cell1', 'cell2'], ...], ...]
    """
    for table_payload in table_payloads:
        if not table_payload:
            continue

        # Count columns from first row
        num_cols = len(table_payload[0]) if table_payload else 1

        # Create the table (without style to ensure compatibility)
        table = doc.add_table(rows=0, cols=num_cols)

        # Populate rows
        for row_data in table_payload:
            row_cells = table.add_row().cells
            for cell_index, cell_value in enumerate(row_data):
                if cell_index < len(row_cells):
                    row_cells[cell_index].text = str(cell_value or "")



def _replace_in_paragraph(paragraph, replacements):
    replaced = 0

    # First try run-level replacement to preserve paragraph/run formatting.
    for run in paragraph.runs:
        run_text = run.text
        updated = run_text
        for tag, value in replacements.items():
            if tag in updated:
                replaced += updated.count(tag)
                updated = updated.replace(tag, value)
        if updated != run_text:
            run.text = updated

    # Fallback for cases where a placeholder spans multiple runs.
    if replaced == 0:
        updated = paragraph.text
        for tag, value in replacements.items():
            if tag in updated:
                replaced += updated.count(tag)
                updated = updated.replace(tag, value)
        if replaced > 0 and updated != paragraph.text:
            paragraph.text = updated

    return replaced


def _replace_in_table(table, replacements):
    replaced = 0
    for row in table.rows:
        for cell in row.cells:
            replaced += _replace_in_container(cell, replacements)
    return replaced


def _replace_in_container(container, replacements):
    replaced = 0
    for paragraph in container.paragraphs:
        replaced += _replace_in_paragraph(paragraph, replacements)

    for table in container.tables:
        replaced += _replace_in_table(table, replacements)

    return replaced


def _copy_section_from_source(source_doc, target_doc, source_title):
    """
    Copy all paragraphs and tables from a section in source_doc to target_doc,
    preserving formatting.
    
    Args:
        source_doc: Source Document object
        target_doc: Target Document object to append to
        source_title: Title text to match (e.g., "1. Introduction")
    """
    if not source_doc or not source_title:
        return
    
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    
    # Normalize the source title for matching
    normalized_source = source_title.strip().lower()
    
    # Find the starting paragraph
    paragraphs = list(source_doc.paragraphs)
    start_index = None
    
    for idx, p in enumerate(paragraphs):
        if p.text.strip().lower() == normalized_source:
            start_index = idx
            break
    
    if start_index is None:
        return
    
    # Copy starting paragraph (the heading itself)
    copy_paragraph_with_formatting(paragraphs[start_index], target_doc)
    
    # Copy subsequent paragraphs until the next section heading
    for idx in range(start_index + 1, len(paragraphs)):
        p = paragraphs[idx]
        
        # Stop if we hit another section heading
        if _is_heading_paragraph(p):
            break
        
        copy_paragraph_with_formatting(p, target_doc)


def _fill_by_section_headings(doc, fields):
    """Fallback fill mode for templates that use section titles instead of tags.
    
    Handles both plain text fields and rich fields with {text, tables} structure."""
    # Extract metadata fields (not meant to be filled into template)
    fields.pop("_source_doc", None)
    fields.pop("_source_sections", None)
    fields.pop("_source_titles", None)
    
    paragraphs = doc.paragraphs
    filled_fields = set()

    for index, paragraph in enumerate(paragraphs):
        heading_key = _normalize_heading(paragraph.text)
        field_key = SECTION_TO_FIELD.get(heading_key)
        if not field_key:
            continue

        field_value = fields.get(field_key, "")

        # Handle rich structure: {text, tables} or plain string
        if isinstance(field_value, dict):
            text = str(field_value.get("text", "") or "")
            tables = field_value.get("tables", []) or []
        else:
            text = str(field_value or "")
            tables = []

        if not text.strip() and not tables:
            continue

        # Fill inside template flow; do not append source sections to document body.
        next_index = index + 1
        if next_index < len(paragraphs):
            next_paragraph = paragraphs[next_index]
            if _is_heading_paragraph(next_paragraph):
                next_paragraph.insert_paragraph_before(text)
            else:
                next_paragraph.text = text
        else:
            if text.strip():
                doc.add_paragraph(text)

        # Insert tables after the section content
        if tables:
            insert_doc_tables(doc, tables)

        filled_fields.add(field_key)


def fill_template(template_path, output_path, fields):
    """
    Fill a DOCX template by replacing predefined tags with provided values.

    Required tags:
    <ProductSummary>, <ValueProposition>, <ProductDescription>,
    <Requirements>, <Scope>, <SLA>, <OperationalSupport>
    """
    doc = Document(template_path)

    replacements = {
        "<ProductSummary>": _extract_text(fields.get("ProductSummary", "")),
        "<ValueProposition>": _extract_text(fields.get("ValueProposition", "")),
        "<ProductDescription>": _extract_text(fields.get("ProductDescription", "")),
        "<Requirements>": _extract_text(fields.get("Requirements", "")),
        "<Scope>": _extract_text(fields.get("Scope", "")),
        "<SLA>": _extract_text(fields.get("SLA", "")),
        "<OperationalSupport>": _extract_text(fields.get("OperationalSupport", "")),
    }

    replaced_count = _replace_in_container(doc, replacements)

    for section in doc.sections:
        replaced_count += _replace_in_container(section.header, replacements)
        replaced_count += _replace_in_container(section.footer, replacements)

    # If template contains no known tags, fill by section headings instead.
    if replaced_count == 0:
        _fill_by_section_headings(doc, fields)

    doc.save(output_path)


def generate_presales(fields, template_path, output_path):
    doc = Document(template_path)

    def replace(tag, value):
        for p in doc.paragraphs:
            if tag in p.text:
                p.text = p.text.replace(tag, value)

    for key, value in fields.items():
        replace(f"<{key}>", value)

    doc.save(output_path)